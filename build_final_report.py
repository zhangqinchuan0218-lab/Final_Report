from copy import deepcopy
from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = Path(r"C:\Users\61461\Desktop\summary2.docx")
OUT = Path(r"F:\esp32_FinalProgram\final_report.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_sep, fld_end])


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p._p.clear_content()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")


def remove_default_empty_paragraph(doc):
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, bold, color in [
        ("Title", 20, True, "1F4E79"),
        ("Heading 1", 14, True, "1F4E79"),
        ("Heading 2", 12, True, "1F4E79"),
        ("Heading 3", 11, True, "000000"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for style_name in ["Header", "Footer"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9)


def setup_section(section):
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_cover(doc, word_count):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GENG4412/5512 Engineering Research Project Part 2")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Software Control System for an Automatic Hair Cutting Machine")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Report")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    lines = [
        "Student name: ______________________________",
        "Student number: ____________________________",
        "Supervisor(s): _____________________________",
        "School of Engineering, University of Western Australia",
        f"Main body word count: {word_count}",
        f"Date of submission: {date.today().strftime('%d %B %Y')}",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    doc.add_page_break()


def add_declaration(doc):
    add_heading(doc, "STUDENT DECLARATION", 1)
    add_heading(doc, "My contribution to the project", 2)
    doc.add_paragraph(
        "The report describes the author's personal contribution to the software design, ESP32 bus-servo communication, multi-motor workflow logic, sensor response, pause and recovery mechanism, and preliminary verification of the automatic hair cutting machine prototype."
    )
    add_heading(doc, "Use of AI tools", 2)
    doc.add_paragraph(
        "AI tools were used to assist with document formatting, structure, and language polishing. The technical content, project work, design decisions, and reported outcomes remain the author's own work and should be checked before submission."
    )
    doc.add_paragraph(
        "I certify that the above information is correct, and that the attached work submitted for assessment is my own work and all material drawn from other sources has been acknowledged and referenced."
    )
    doc.add_paragraph("Student signature ________________________________    Date ________________")
    doc.add_paragraph("Supervisor signature _____________________________    Date ________________")
    doc.add_page_break()


def clean_heading_text(text):
    text = re.sub(r"\s+", " ", text.strip())
    text = re.sub(r"^(\d+(?:\.\s*\d+)*\.?)\s*", lambda m: m.group(1).replace(" ", "") + " ", text)
    fixes = {
        "3.1.1 background and motivation": "3.1.1 Background and Motivation",
        "3.1.1background and motivation": "3.1.1 Background and Motivation",
        "3.1.2 problem statement": "3.1.2 Problem Statement",
        "3.2 Literature review": "3.2 Literature Review",
        "3.3 project objectives": "3.3 Project Objectives",
        "4.4 Alternative generation process and scheme comparison": "4.4 Alternative Generation Process and Scheme Comparison",
        "4.6. 2. Lap position feedback and unit position definition": "4.6.2 Lap Position Feedback and Unit Position Definition",
        "4.6. 3. Cross-lap problem analysis": "4.6.3 Cross-Lap Problem Analysis",
        "4.6. 4. Cross-circle correction method": "4.6.4 Cross-Circle Correction Method",
        "4.6. 6. Establishment of preset lap targets": "4.6.6 Establishment of Preset Lap Targets",
        "4.6. 7. PID control algorithm and speed closed-loop correction": "4.6.7 PID Control Algorithm and Speed Closed-Loop Correction",
    }
    return fixes.get(text, text)


def heading_level(text):
    m = re.match(r"^(\d+(?:\.\d+)*)\.?\s+", text)
    if not m:
        return None
    depth = len(m.group(1).split("."))
    if depth <= 1:
        return 1
    if depth == 2:
        return 2
    return 3


def is_equation_like(text):
    if len(text) > 80:
        return False
    symbols = ["∈", "Δ", "=", "4096", "360°", "overcurrent", "timeout"]
    return any(s in text for s in symbols) and not re.match(r"^\d+(?:\.\d+)*", text)


def add_front_matter(doc, summary_text, toc_items):
    add_heading(doc, "Project Summary", 1)
    doc.add_paragraph(summary_text)
    doc.add_page_break()

    add_heading(doc, "Acknowledgements", 1)
    doc.add_paragraph("Acknowledgements to be completed before submission.")
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    for level, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 * (level - 1))
        p.paragraph_format.space_after = Pt(2)
        p.add_run(title)
    doc.add_page_break()

    add_heading(doc, "List of Tables", 1)
    doc.add_paragraph("Table 1: Software verification outcomes.")
    doc.add_page_break()

    add_heading(doc, "Nomenclature", 1)
    terms = [
        ("ESP32", "Microcontroller used for prototype control"),
        ("M1-M5", "Bus servo motors used in the prototype"),
        ("PID/PD", "Proportional-integral-derivative / proportional-derivative control concepts"),
        ("SC09", "Bus servo communication interface"),
        ("ST3215-HS", "Bus servo actuator used in the prototype"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Symbol / term"
    table.rows[0].cells[1].text = "Meaning"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for term, meaning in terms:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = meaning
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    doc.add_page_break()


def copy_table(doc, src_table, caption):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    table.style = "Table Grid"
    for i, row in enumerate(src_table.rows):
        for j, cell in enumerate(row.cells):
            target = table.cell(i, j)
            target.text = cell.text
            target.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(target)
            for p in target.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                if j == len(row.cells) - 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(target, "D9EAF7")
                for p in target.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def build():
    src = Document(SRC)
    summary_text = ""
    for p in src.paragraphs:
        t = " ".join(p.text.split())
        if t and t.lower() != "summary":
            summary_text = t
            break

    main_text = "\n".join(
        p.text for p in src.paragraphs if p.text.strip() and p.text.strip().lower() != "summary"
    )
    word_count = len(re.findall(r"\b[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?\b", main_text))
    toc_items = []
    for p in src.paragraphs:
        text = " ".join(p.text.split())
        if not text:
            continue
        cleaned = clean_heading_text(text)
        level = heading_level(cleaned)
        if level:
            toc_items.append((min(level, 3), cleaned))

    doc = Document()
    remove_default_empty_paragraph(doc)
    configure_styles(doc)
    setup_section(doc.sections[0])
    add_page_number(doc.sections[0])

    add_cover(doc, word_count)
    add_declaration(doc)
    add_front_matter(doc, summary_text, toc_items)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_section(body_section)
    body_section.footer.is_linked_to_previous = True

    table_inserted = False
    for idx, p in enumerate(src.paragraphs):
        text = " ".join(p.text.split())
        if not text or text.lower() == "summary":
            continue
        if text == summary_text:
            continue

        cleaned = clean_heading_text(text)
        level = heading_level(cleaned)
        if level:
            add_heading(doc, cleaned, min(level, 3))
        elif is_equation_like(cleaned):
            eq = doc.add_paragraph()
            eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = eq.add_run(cleaned)
            run.italic = True
        else:
            para = doc.add_paragraph(cleaned)
            para.paragraph_format.first_line_indent = Cm(0)

        if not table_inserted and idx > 148 and src.tables:
            copy_table(doc, src.tables[0], "Table 1: Software verification outcomes.")
            table_inserted = True

    add_heading(doc, "References", 1)
    doc.add_paragraph(
        "Reference list to be completed before submission. Ensure all literature, standards, datasets, figures, tables, and copied or adapted technical material cited in the report are included here in a consistent referencing style."
    )
    add_heading(doc, "Appendices", 1)
    doc.add_paragraph("Appendices may include source code extracts, raw test logs, wiring diagrams, and additional verification data.")

    doc.core_properties.title = "Software Control System for an Automatic Hair Cutting Machine"
    doc.core_properties.subject = "GENG4412/5512 Final Report"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
