from pathlib import Path
from docx import Document

summary = Path(r"C:\Users\61461\Desktop\summary2.docx")
template = Path(r"C:\Users\61461\Downloads\GENG4412_GENG5512 Final Report Template S1 2026.docx")

for label, path in [("SUMMARY", summary), ("TEMPLATE", template)]:
    doc = Document(path)
    print(f"\n== {label} ==")
    print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "sections", len(doc.sections))
    print("styles:", ", ".join(s.name for s in doc.styles if s.type in (1, 2))[:1500])
    print("\nfirst paragraphs:")
    for i, p in enumerate(doc.paragraphs[:80], 1):
        text = " ".join(p.text.split())
        if text:
            print(f"{i:03d} [{p.style.name}] {text[:220]}")
    print("\ntables:")
    for ti, table in enumerate(doc.tables[:5], 1):
        print(f"table {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for row in table.rows[:3]:
            print(" | ".join(" ".join(c.text.split())[:80] for c in row.cells))
