from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches


ROOT = Path(r"F:\esp32_FinalProgram")
SUMMARY = Path(r"F:\5512\摘要summary.docx")
TEMPLATE = Path(r"C:\Users\61461\Downloads\GENG4412_GENG5512 Final Report Template S1 2026.docx")
OUT = Path(r"F:\5512\Final_Report_完善版.docx")
CODE = ROOT / "Final" / "Final.ino"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    except Exception:
        pass


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.15


def add_equation(doc: Document, text: str, number: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{text}    ({number})")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def existing_paragraphs() -> list[str]:
    src = Document(str(SUMMARY))
    return [p.text.strip() for p in src.paragraphs if p.text.strip()]


def code_value(pattern: str, default: str = "") -> str:
    text = CODE.read_text(encoding="utf-8", errors="ignore")
    m = re.search(pattern, text)
    return m.group(1) if m else default


def main() -> None:
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    set_margins(doc)

    add_title(doc, "ESP32-Based Multi-Motor Control System for an Automatic Haircutting Prototype")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("GENG4412/GENG5512 Engineering Research Project Final Report").bold = True
    doc.add_paragraph("Student: ____________________    Student number: ____________________")
    doc.add_paragraph("Supervisor: ____________________")
    doc.add_paragraph("School of Engineering, The University of Western Australia")
    doc.add_paragraph("Word count: ____________________")
    doc.add_paragraph("Date: ____________________")
    doc.add_page_break()

    add_heading(doc, "Student Declaration", 1)
    add_heading(doc, "My contribution to the project", 2)
    add_para(doc, "本项目为团队设计与制作项目。机械结构建模、3D 打印和部分实体装配主要由团队其他成员完成。本人主要完成 ESP32 与总线舵机的通信、M1 至 M5 多电机控制程序、轮子模式多圈相对位置控制、状态机动作流程、M2 回零与安全保护、压力传感器响应、暂停恢复逻辑以及串口调试信息设计。")
    add_heading(doc, "Use of AI tools", 2)
    add_para(doc, "本文档在整理表达、检查结构完整性和排版方面使用了 AI 工具辅助，但项目设计、程序调试、实验观察和技术判断由本人完成。所有程序参数、公式和结果讨论均基于本人项目中的实际控制程序与调试记录。")
    doc.add_page_break()

    add_heading(doc, "Project Summary", 1)
    paras = existing_paragraphs()
    for text in paras[:3]:
        if "摘要" in text.lower():
            continue
        add_para(doc, text)
    doc.add_page_break()

    add_heading(doc, "Nomenclature", 1)
    add_table(doc, ["Symbol", "Meaning", "Unit or note"], [
        ("N", "编码器累计位置", "counts"),
        ("N0", "动作开始时累计位置", "counts"),
        ("Nt", "目标累计位置", "counts"),
        ("e", "位置误差，e = Nt - N", "counts"),
        ("v", "由累计位置差分得到的速度", "counts/s"),
        ("u", "发送给轮子模式电机的速度指令", "servo command"),
        ("Kp", "位置误差增益", "-"),
        ("Kv", "速度反馈增益", "-"),
        ("I", "舵机反馈电流读数", "raw count,约 6.5 mA/count"),
        ("L", "舵机负载反馈读数", "raw count"),
    ])
    doc.add_page_break()

    for text in paras[3:]:
        heading_match = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)", text)
        if heading_match and len(text) < 80:
            level = min(3, heading_match.group(1).count(".") + 1)
            add_heading(doc, text, level)
        else:
            add_para(doc, text)

    add_heading(doc, "4.3 控制模型与关键公式", 2)
    add_para(doc, "为了满足 rubric 中关于方法可复现性的要求，本节补充程序中最核心的控制公式。ST3215 舵机单圈反馈范围为 0 至 4095，因此程序将一圈定义为 4096 个编码单位。轮子模式本身只接收速度指令，不能直接设定多圈绝对位置，所以本项目在 ESP32 程序中自行建立累计位置，并在外部控制器中形成相对位置闭环。")
    add_equation(doc, "N = 4096 n + p", "4.1")
    add_para(doc, "式中，N 为连续累计位置，n 为通过跨圈检测得到的圈数，p 为舵机单圈位置反馈。当相邻两次读数跨过 0/4095 边界时，程序判断发生跨圈跳变，并对累计圈数进行加减。这样，单圈反馈被转换为可连续增加或减少的相对位置。")
    add_equation(doc, "e(k) = N_t - N(k)", "4.2")
    add_para(doc, "式中，e(k) 为第 k 个控制周期的位置误差，Nt 为目标累计位置。每次动作开始时，目标位置由当前位置和目标相对位移相加得到。")
    add_equation(doc, "v(k) = [N(k) - N(k-1)] / Δt", "4.3")
    add_para(doc, "速度 v(k) 由两个控制周期之间的累计位置变化估算。程序的电机控制任务约每 20 ms 运行一次，因此该速度反馈能够反映电机是否正在靠近目标，也用于堵转判断。")
    add_equation(doc, "u_raw(k) = Kp e(k) - Kv v(k)", "4.4")
    add_para(doc, "式中，u_raw(k) 为未限幅的轮子模式速度指令。Kp 控制电机向目标位置靠近的强度，Kv 抑制接近目标时的速度，从而减少过冲。程序中 M1、M2、M3 分别使用不同的 Kp 和 Kv，以适应不同负载和机构方向。")
    add_equation(doc, "u(k)=sat(u_raw(k), -u_max, u_max)", "4.5")
    add_para(doc, "由于舵机和 3D 打印结构不能承受过大冲击，速度指令会被限制在最大速度范围内。对于 M2，程序还加入了方向最小速度、重力补偿和速度斜坡限制，使竖直方向的启动更可靠，同时避免突然加速。")
    add_equation(doc, "u(k)=u(k-1)+clip[u_raw(k)-u(k-1), -Δu_max, Δu_max]", "4.6")
    add_para(doc, "式 (4.6) 表示速度斜坡限制。该限制防止 M2 指令从 0 突然跳变到较大速度，能够降低机械冲击，也能减少由于瞬间负载变化导致的误判堵转。")

    add_heading(doc, "4.4 关键行程参数推导", 2)
    add_para(doc, "程序中所有主要行程都由 ONE_TURN = 4096 推导而来。这样做的好处是参数与舵机编码器分辨率直接对应，后续若需要改变步距，只需修改分数关系即可。")
    add_table(doc, ["Parameter", "Program expression", "Derived value", "Function"], [
        ("M1_STEP", "-ONE_TURN / 4", "-1024 counts", "M1 每层横向覆盖后旋转四分之一圈"),
        ("M1_TOTAL", "-ONE_TURN * 7", "-28672 counts", "M1 总覆盖范围"),
        ("M2_INITIAL_LIFT", "-ONE_TURN / 8", "-512 counts", "M2 回零后先抬升，离开底部机械接触区"),
        ("M2_STEP", "-ONE_TURN / 20", "-204 counts", "M2 每次纵向分层进给"),
        ("M2_LAYER_TOTAL", "ONE_TURN * 3", "12288 counts", "M2 每一层累计三圈后触发 M1"),
        ("M3_SMALL_RETRACT_DELTA", "-ONE_TURN / 20", "-204 counts", "压力触发后 M3 小幅回退"),
        ("M3_APPROACH_MAX_TRAVEL", "ONE_TURN * 2", "8192 counts", "M3 接近动作最大安全行程"),
    ])
    add_para(doc, "需要注意的是，C++ 中整数除法会舍去小数。例如 4096/20 的数学值为 204.8，而程序中的整数结果为 204。因此，M2 每次分层的实际目标为 204 counts，三圈累计目标为 12288 counts。该误差对单步很小，但在长期重复动作中应通过实际测试校准。")

    add_heading(doc, "4.5 M2 回零与安全判据", 2)
    add_para(doc, "M2 是竖直方向电机，回零过程直接影响后续所有分层动作。程序先让 M2 以较低速度向下运动，并持续读取电流、负载、速度和累计行程。底部并不是单纯依靠速度为零判断，而是要求电流或负载出现底部信号，并在短时间内保持接近停止状态。")
    add_equation(doc, "H = |N - N_start|", "4.7")
    add_equation(doc, "ΔI = |I| - |I_start|,    ΔL = |L| - |L_start|", "4.8")
    add_para(doc, "式中，H 为回零期间已经走过的行程，ΔI 和 ΔL 分别为相对起始状态的电流和负载变化。当电流读数达到底部信号阈值，或负载变化达到阈值，同时速度足够低并持续确认一段时间后，程序才将当前位置设为 M2 的逻辑零点。")
    add_para(doc, "根据调试记录，ST 系列舵机反馈的 Current 不是直接以 mA 显示，而是舵机协议给出的原始计数值。按本项目调试采用的换算估计，1 count 约等于 6.5 mA，因此 77 count 约为 500 mA。程序中 M2 回零电流保护设置为 77，底部触发信号设置为 10 count；这样底部触发较敏感，而真正急停保护更保守。")

    add_heading(doc, "4.6 系统状态机流程", 2)
    add_para(doc, "整个控制流程采用状态机实现。启动后，系统依次进入 M2 回零、M2 离底初始抬升、M5 保持准备、剪发动作循环和 M1 层间移动。一次剪发动作由 M2 上升、M3 接近、压力触发停止、M3 小回退、M4 横向动作、M5 剪切、M3 回原位组成。该流程避免在 loop() 中使用长时间 delay()，因此按钮、压力传感器和异常保护可以在运行过程中持续被检测。")
    add_table(doc, ["State/action", "Purpose", "Transition condition"], [
        ("M2_HOMING_DOWN", "寻找 M2 底部并建立逻辑零点", "底部信号确认或异常保护触发"),
        ("M2_INITIAL_LIFT", "回零后上升 1/8 圈，离开底部", "M2 到达目标"),
        ("PREPARE_M5", "让 M5 保持 820 的准备位置", "位置到达或准备超时后继续"),
        ("M2_STEP", "M2 进行一次纵向分层进给", "M2 到位"),
        ("M3_APPROACH", "M3 慢速靠近，等待压力触发", "压力触发或最大行程/超时"),
        ("M4/M5 actions", "完成横向捕捉和剪切动作", "位置到达"),
        ("M1_STEP", "M2 累计三圈后 M1 转动 1/4 圈", "M1 到位后进入下一层"),
    ])

    add_heading(doc, "5. Results and Discussion", 1)
    add_para(doc, "本项目的主要结果不是单个电机的孤立转动，而是一套可以在 ESP32 上运行的多电机控制框架。通过轮子模式累计位置控制，M1、M2 和 M3 能够突破常规位置模式的一圈限制，实现相对多圈运动；通过位置模式控制，M4 和 M5 能够在规定范围内进行重复定位。实际调试表明，系统能够完成 M2 分层、M3 接近、M4 横向动作、M5 剪切和 M1 层间旋转的顺序联动。")
    add_para(doc, "在调试过程中，M2 是最关键也最容易出现问题的电机。由于 M2 承受竖直方向负载，回零和上升阶段会出现电流升高、速度为零和机械卡顿。早期程序曾出现尚未到达底部就误判回零，或者刚回零后立即上升时因底部机械接触导致堵转。最终程序通过降低回零速度、加入底部电流信号、加入最小回零行程、设置回零确认时间、回零后先上升 1/8 圈以及更谨慎的电流保护，显著提高了 M2 启动和分层的稳定性。")
    add_para(doc, "M3 的结果表明，压力传感器触发逻辑对安全性非常重要。M3 接近头部方向时，如果只依靠固定行程，可能在不同头部距离或机械装配误差下产生过压风险。当前程序将压力信号作为停止条件，并在压力触发后记录 touch 位置，再执行小幅回退。这样可以让接近动作根据实际接触情况停止，而不是完全依赖预设位置。")
    add_para(doc, "系统仍然存在明显限制。第一，舵机反馈电流和负载并不是高精度测量值，只适合作为保护和趋势判断，不能替代独立电流传感器。第二，软件急停依赖程序循环和通信正常工作，不能替代硬件急停。第三，M2 和 M3 的负载会随机构位置变化，因此 PID 参数只能在当前机械状态下取得较好折中。第四，由于多个舵机共用串口总线，偶尔会出现反馈读数为 -1 的通信失败，需要通过重试和容错逻辑降低影响。")

    add_heading(doc, "6. Conclusions and Future Work", 1)
    add_para(doc, "本项目完成了一套基于 ESP32 的自动理发机多电机控制系统。该系统能够同时管理 ST3215 和 SC09 总线舵机，支持位置模式和轮子模式混合控制，并通过软件累计位置实现 M1、M2 和 M3 的相对多圈控制。程序采用状态机和 FreeRTOS 多任务结构，使运动控制、按钮输入、压力传感器检测和串口调试能够并行运行。")
    add_para(doc, "项目最重要的工程贡献在于将原本分散的单电机测试扩展为完整的五电机动作流程，并针对 M2 回零、M3 接近和系统急停等实际问题建立了可调试的保护逻辑。通过参数调试，系统能够更稳定地完成回零、分层、靠近、横移、剪切和复位动作，为后续自动理发机的机械集成和真实剪发测试提供了软件基础。")
    add_para(doc, "未来工作应首先加入硬件急停、限位开关和独立电流检测，使安全保护不完全依赖软件和舵机反馈。其次，应对 M2 和 M3 建立更系统的负载测试，记录不同位置、不同速度和不同电流限制下的成功率，从而进一步优化 Kp、Kv、速度斜坡和电流阈值。第三，应将机械结构、剪发工具和传感器进行完整集成，并在安全替代物上进行重复剪切实验，量化覆盖范围、定位误差、动作周期和异常停止响应时间。")

    add_heading(doc, "References", 1)
    add_para(doc, "Espressif Systems. ESP32 technical reference manual and programming documentation.")
    add_para(doc, "Waveshare. ST3215-HS serial bus servo user manual and communication protocol documentation.")
    add_para(doc, "Waveshare. SC09 serial bus servo user manual and communication protocol documentation.")
    add_para(doc, "FreeRTOS. Real-time kernel documentation for task scheduling and embedded control.")

    add_heading(doc, "Appendix A: Key Program Configuration", 1)
    add_para(doc, "本附录列出最终程序中与控制效果最直接相关的参数，便于后续复现实验和继续调试。")
    add_table(doc, ["Group", "Parameter", "Value in program", "Meaning"], [
        ("Encoder", "ONE_TURN", "4096", "one mechanical revolution in encoder counts"),
        ("M2", "M2_HOMING_DOWN_SPEED", "150", "slow downward homing speed"),
        ("M2", "M2_HOMING_DOWN_PROBE_SPEED", "220", "probe speed if no movement is detected"),
        ("M2", "M2_HOMING_CURRENT_LIMIT", "77", "M2 current protection during homing and lifting"),
        ("M2", "M2_HOMING_BOTTOM_CURRENT", "10", "current signal used for bottom detection"),
        ("M2", "M2_STALL_GRACE_MS", "3000", "time window before stall protection becomes active"),
        ("M3", "M3_APPROACH_SPEED", "120", "slow approach speed"),
        ("M4", "M4_CCW_POS / M4_CW_POS", "568 / 3528", "calibrated lateral motion limits"),
        ("M5", "M5_HOLD_POS / M5_CUT_POS", "820 / 700", "open and cutting positions"),
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
