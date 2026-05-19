from pathlib import Path
from docx import Document
from pypdf import PdfReader


ROOT = Path(r"F:\esp32_FinalProgram")
OUT = ROOT / "report_work"
OUT.mkdir(exist_ok=True)

paths = {
    "summary": Path(r"F:\5512\摘要summary.docx"),
    "template": Path(r"C:\Users\61461\Downloads\GENG4412_GENG5512 Final Report Template S1 2026.docx"),
    "rubric": Path(r"C:\Users\61461\Downloads\GENG4412_GENG5512 Final Report Rubric S1 2026 - Tagged.pdf"),
    "code": ROOT / "Final" / "Final.ino",
}


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    lines = []
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if text:
            lines.append(f"P{i}: {text}")
    for ti, table in enumerate(doc.tables, 1):
        lines.append(f"[TABLE {ti}]")
        for row in table.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        pages.append(f"--- PAGE {i} ---\n{text}")
    return "\n\n".join(pages)


(OUT / "summary_text.txt").write_text(docx_text(paths["summary"]), encoding="utf-8")
(OUT / "template_text.txt").write_text(docx_text(paths["template"]), encoding="utf-8")
(OUT / "rubric_text.txt").write_text(pdf_text(paths["rubric"]), encoding="utf-8")
(OUT / "code_text.txt").write_text(paths["code"].read_text(encoding="utf-8", errors="ignore"), encoding="utf-8")

print(OUT)
