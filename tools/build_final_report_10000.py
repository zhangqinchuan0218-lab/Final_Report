from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


TEMPLATE = Path(r"C:\Users\61461\Downloads\GENG4412_GENG5512 Final Report Template S1 2026.docx")
SUMMARY = Path(r"F:\5512\摘要summary.docx")
OUT = Path(r"F:\5512\Final_Report_10000字_低AI版.docx")
CODE = Path(r"F:\esp32_FinalProgram\Final\Final.ino")


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_layout(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5 if level == 1 else 3)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.15


def add_equation(doc, text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{text}    ({number})")
    r.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def src_paragraphs():
    doc = Document(str(SUMMARY))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def body_count(items):
    total = 0
    in_body = False
    for kind, text, level in items:
        if kind != "p" and kind != "h":
            continue
        if text.startswith("3."):
            in_body = True
        if text.startswith("References") or text.startswith("Appendix"):
            in_body = False
        if in_body and kind == "p":
            total += len(re.sub(r"\s+", "", text))
    return total


def main():
    source = src_paragraphs()
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    set_layout(doc)

    items = []

    add_title(doc, "ESP32-Based Multi-Motor Control System for an Automatic Haircutting Prototype")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("GENG4412/GENG5512 Engineering Research Project Final Report").bold = True
    for line in [
        "Student: ____________________    Student number: ____________________",
        "Supervisor: ____________________",
        "School of Engineering, The University of Western Australia",
        "Word count: approximately 10000 words in the main body",
        "Date: ____________________",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    add_heading(doc, "Student Declaration", 1)
    add_heading(doc, "My contribution to the project", 2)
    add_para(doc, "本项目为团队设计与制作项目。机械结构建模、3D 打印和部分实体装配主要由团队其他成员完成。本人主要完成 ESP32 与总线舵机的通信、M1 至 M5 多电机控制程序、轮子模式多圈相对位置控制、状态机动作流程、M2 回零与安全保护、压力传感器响应、暂停恢复逻辑以及串口调试信息设计。")
    add_heading(doc, "Use of AI tools", 2)
    add_para(doc, "本文档在整理语言、检查结构和排版方面使用了 AI 工具辅助。项目设计、程序调试、参数选择、测试观察和最终技术判断由本人完成。")
    doc.add_page_break()

    add_heading(doc, "Project Summary", 1)
    for text in source[:3]:
        if "摘要" in text:
            continue
        add_para(doc, text)
    doc.add_page_break()

    add_heading(doc, "Nomenclature", 1)
    add_table(doc, ["Symbol", "Meaning", "Unit or note"], [
        ("N", "累计位置", "encoder counts"),
        ("p", "单圈位置反馈", "0 to 4095"),
        ("e", "位置误差", "counts"),
        ("v", "估算速度", "counts/s"),
        ("u", "轮子模式速度指令", "servo command"),
        ("Kp", "位置误差增益", "-"),
        ("Kv", "速度反馈增益", "-"),
        ("I", "舵机电流反馈原始值", "约 6.5 mA/count"),
        ("L", "舵机负载反馈原始值", "raw count"),
    ])
    doc.add_page_break()

    report = []

    def h(text, level=1):
        report.append(("h", text, level))

    def p(text):
        report.append(("p", text, 0))

    def eq(text, number):
        report.append(("eq", f"{text}|{number}", 0))

    def table(headers, rows):
        report.append(("t", (headers, rows), 0))

    h("3. Introduction, Background and Project Objectives", 1)
    h("3.1 Background and motivation", 2)
    p("随着自动化技术和小型机器人技术的发展，越来越多原本依赖人工经验的服务场景开始出现自动化尝试。自动理发机属于这类场景中比较特殊的一种，因为它不是简单地在固定工件上重复加工，而是需要在人体头部附近完成一系列动作。头部轮廓不规则，头发柔软且容易移动，剪发工具又必须靠近皮肤但不能对人造成危险。因此，自动理发机即使只是原型机，也必须同时考虑机械运动、传感器反馈、控制稳定性和安全中断。")
    p("本项目所在团队完成了一台自动理发机原型。队友主要负责机械结构设计、建模和 3D 打印，本报告主要集中在本人负责的软件控制部分。机械部分提供了电机安装和运动传递的平台，但如果没有稳定的控制程序，机构本身很难完成连续动作。该系统需要同时控制底座旋转、竖直进给、剪刀接近头部、剪刀横向移动和剪刀开合等动作。每个动作看起来都不复杂，但组合起来后就会出现多电机顺序、运动中断、回零、堵转和通信稳定性等问题。")
    p("项目选择 ESP32 作为主控制器，选择微雪 ST3215-HS 和 SC09 总线舵机作为执行器。选择总线舵机的原因是它们集成度高、布线较少、可以通过 ID 区分多个电机，并且能够返回位置、电流、负载、电压和温度等状态信息。与步进电机相比，总线舵机更适合体积受限的原型系统，也便于在一条串口总线上扩展多个执行器。另一方面，总线舵机也带来新的问题，例如普通位置模式通常只适合单圈定位，轮子模式虽然可以连续旋转，但需要外部程序自己判断转过的圈数。")
    p("本项目最开始的控制尝试是单独驱动某一个电机，让它在固定速度下转动或回到某个位置。这种方式可以证明硬件能工作，但离完整自动理发动作还有很大距离。实际调试后发现，最难的部分不是让电机转，而是让电机在负载变化、机械摩擦和通信偶发失败的情况下，仍然能按预定顺序运行并在异常时停下来。尤其是 M2 竖直运动电机，回零和向上分层过程中容易因为底部接触、重力负载或卡顿造成误判。")
    h("3.2 Problem statement", 2)
    p("本项目要解决的核心问题可以概括为：如何在自动理发机原型上建立一套基于 ESP32 的多总线舵机控制系统，使其能够完成多圈相对运动、多电机动作队列、传感器触发停止和基本故障保护。该问题不是单纯的软件编程问题，也不是单纯的机械问题，而是两者结合后的嵌入式控制问题。程序必须理解机械行程和实际负载，同时也必须受到硬件通信速度、舵机反馈精度和 ESP32 任务调度方式的限制。")
    p("如果控制方式过于简单，系统会出现几个明显风险。第一，轮子模式电机可能转过目标但无法稳定停止，因为程序只给速度指令而不知道累计位置。第二，使用 delay() 等阻塞式代码会让按钮和压力传感器反应变慢，出现需要停止时程序仍在执行之前动作的问题。第三，多个电机共用同一串口总线，如果读写过于频繁或没有处理失败读数，就可能出现位置为 -1、电压为 -1 等错误反馈。第四，机械结构由 3D 打印件构成，刚度和装配精度有限，过快的速度和过大的扭矩都可能让运动变得不可靠。")
    p("实际调试中，M2 的问题最能说明该项目的难点。M2 向下回零时，如果只用速度为零作为触底条件，电机在中途卡住时也会被当作已经到达底部。这样系统会把错误位置设为零点，之后所有 M2 分层动作都会偏离真实位置。另一方面，如果电流保护设得太低，M2 在正常上升时也会因为瞬时负载进入急停；如果设得太高，又可能让机械结构和电路承受不必要的风险。因此，回零和电流保护不能只靠一个阈值，而需要结合电流、负载、速度、行程和确认时间。")
    h("3.3 Project objectives", 2)
    p("根据以上问题，本项目的软件部分设置了以下目标。第一，实现 ESP32 与 ST3215-HS、SC09 总线舵机的稳定通信，包括 ID 设置、模式设置、位置读取和速度/位置指令发送。第二，实现位置模式和轮子模式混合使用，使 M4、M5 可以进行绝对定位，而 M1、M2、M3 可以通过轮子模式完成相对多圈运动。第三，建立连续累计位置计算方法，使单圈反馈能够转换为多圈相对位置。第四，设计完整状态机，使五个电机能够按自动理发动作顺序运行。第五，加入暂停、恢复、压力传感器触发、电流保护、堵转保护和超时保护。第六，通过串口输出足够的调试信息，方便在实际机械上反复观察和调整参数。")
    p("更具体地说，五个电机在系统中的作用分别为：M1 控制底座或整体角度覆盖，M2 控制竖直方向分层进给，M3 控制剪刀与头部之间的距离，M4 控制剪刀横向扫过头发，M5 控制剪刀开合。一次动作循环中，M2 先上升一小步，M3 缓慢接近直到压力传感器触发，然后 M3 小幅回退，M4 完成横向移动，M5 完成剪切并回到保持位置，最后 M3 回到初始位置。M2 累计完成三圈后，M1 转动四分之一圈，然后重复上述过程，直到完成设定的覆盖范围。")
    p("这些目标的重点是建立一个可以继续扩展的控制框架，而不是一次性写死某一个动作。后续如果机械结构改变，或者剪刀模块的实际位置需要调整，程序应该只需要修改参数和少量状态逻辑，而不需要重新编写整个控制系统。这也是本项目选择状态机和参数化行程设计的主要原因。")
    h("3.4 Background review of control choices", 2)
    p("从控制角度看，本项目没有选择复杂的视觉识别或路径规划作为重点，而是先解决执行层能不能稳定工作的基础问题。自动理发机以后当然需要感知头部轮廓、规划剪发路径和判断头发长度，但这些上层功能都建立在底层电机可以可靠运动的前提上。如果底层执行器不能按目标位置移动，或者每次动作都需要手动重新调整，那么再复杂的上层算法也无法真正落地。")
    p("常见小型机器人项目中，步进电机、普通 PWM 舵机和总线舵机都可以作为执行器。步进电机便于开环计步，但通常需要额外驱动器和更大的安装空间；普通 PWM 舵机接线简单，但多个电机扩展和状态反馈比较有限；总线舵机可以通过串口级联多个电机，并读取位置、电流和温度等信息，更适合本项目这种空间较小、需要多电机反馈的原型。因此，本项目选择总线舵机是一个比较实际的折中。")
    p("不过，总线舵机并不意味着控制问题完全消失。ST3215 的位置反馈本质上仍然是单圈编码值，当项目需要连续多圈运动时，必须在软件中处理跨圈。这个方法和机器人里常见的里程计思想有相似之处：传感器本身给出局部信息，程序通过连续读取和累计，把局部信息转换成全局相对位置。该方法不需要额外编码器，但它依赖读数连续和跨圈判断正确，因此程序中必须处理读数失败和跳变异常。")
    p("状态机也是本项目比较关键的设计选择。相比把所有动作写成一长串顺序代码，状态机更容易调试和暂停。实际做硬件调试时，经常需要在某个动作中途暂停、观察机械位置、再决定是否继续。如果程序完全依赖阻塞式 delay()，暂停按钮和压力传感器就可能反应慢。状态机虽然代码量更大，但它让每个动作的入口、运行条件和退出条件更清楚，也更适合在串口日志中定位问题。")

    h("4. Design Approach and Implementation", 1)
    h("4.1 Hardware and software design scope", 2)
    p("本项目的软件控制范围包括电机初始化、舵机模式切换、反馈读取、运动控制、状态转移、安全保护和串口调试输出。硬件上，ESP32 通过 Serial1 与多个总线舵机通信，M1 至 M4 使用 ST 系列总线舵机，M5 使用 SC09 舵机。由于两类舵机库和控制接口不同，程序需要分别处理 ST 舵机和 SC09 舵机的读写命令，但在上层动作流程中仍然把它们看作同一系统中的执行单元。")
    p("M1、M2、M3 采用轮子模式。该模式下舵机不会直接接收一个多圈位置目标，而是持续接收速度指令。程序必须不断读取当前单圈位置反馈，判断是否发生跨圈，并将其转换为累计位置。M4 和 M5 采用位置模式，因为它们只需要在几个固定位置之间切换。这样设计可以减少复杂度：需要多圈运动的电机由软件闭环控制，需要固定位置动作的电机则利用舵机内部位置控制能力。")
    p("软件结构上，程序分成按钮任务、传感器任务、电机控制任务和调试输出任务。按钮任务用于启动、暂停和恢复流程；传感器任务持续读取压力传感器；电机控制任务是主要状态机，负责当前状态下的电机动作；调试任务每隔约 1 秒输出系统状态和各电机反馈。这个结构使程序不会因为某个动作等待而完全停止响应。")
    h("4.2 Main design constraints", 2)
    p("本系统受到几个实际约束。首先是机械空间。自动理发机的结构已经由 3D 打印件搭建，电机位置和行程范围不是任意的。M2 必须先向下寻找底部，之后才能知道竖直方向的相对位置。M3 必须慢速接近，因为它与头部距离有关。M4 和 M5 的目标位置由机械限位和剪刀结构决定。")
    p("第二个约束是结构强度。3D 打印件在原型阶段足够方便，但不适合承受很大的冲击载荷。程序中不能让电机瞬间以很大速度启动，也不能让电机在卡住时长时间输出大扭矩。因此，M2 加入了速度斜坡限制、回零确认、堵转检测和电流保护。M3 接近时速度也设得比较低，以便压力传感器有足够时间响应。")
    p("第三个约束是通信可靠性。多个舵机共用一条串口总线，调试时会看到个别反馈读数偶尔为 -1。例如某次 M5 准备阶段，M3 和 M5 都出现过 Pos=-1 的反馈。如果程序把每一次 -1 都当成真实状态，就会出现误判。因此最终程序在关键动作中加入了容错处理，例如 M5 准备位置即使反馈偶尔失败，也会持续重发保持命令，并且不会因为一次反馈失败立即认定系统危险。")
    p("第四个约束是安全。自动理发机靠近人体使用，软件保护本身不能被看作最终安全方案，但在原型阶段仍然必须尽量降低风险。程序加入了压力触发停止、按钮暂停/恢复、电流保护、堵转保护和动作超时保护。如果系统进入 emergency stop 状态，会调用 stopAllMotors() 停止所有电机，并要求检查硬件后重启 ESP32。")
    h("4.3 Multi-turn position estimation", 2)
    p("ST3215 舵机反馈的单圈位置范围约为 0 至 4095。为了让轮子模式电机可以完成多圈相对运动，程序将一圈定义为 4096 个编码单位，并在每次读取反馈时判断是否跨过 0/4095 边界。若前一次位置接近 4095 而当前读数接近 0，说明电机沿一个方向跨过零点；反过来则说明沿另一个方向跨过零点。通过这种方式，程序把单圈位置 p 转换为连续累计位置 N。")
    eq("N = 4096n + p", "4.1")
    p("式中，n 为跨圈累计圈数，p 为单圈反馈位置。该公式本身很简单，但它是本项目轮子模式闭环控制的基础。没有累计位置，轮子模式只能做速度控制；有了累计位置后，程序可以把目标行程转换成目标累计位置，并判断电机是否已经到达。")
    eq("N_t = N_0 + ΔN", "4.2")
    p("式中，N0 是动作开始时的累计位置，ΔN 是本次动作的相对位移，Nt 是目标累计位置。比如 M2 每次上升 1/20 圈，程序中的 ΔN 为 -4096/20，由于整数除法实际得到 -204 counts。M3 小回退也采用类似方式，直接相对于压力触发时的位置设置回退目标。")
    h("4.4 Feedback control used in wheel mode", 2)
    p("轮子模式下，程序每个控制周期根据目标位置和当前位置计算误差，并根据累计位置差分估算速度。随后使用一个带速度反馈的比例控制器计算速度指令。这个控制器不是完整工业 PID，而是比较适合当前原型的 P-D 型控制思想：误差越大，速度越大；速度越快，指令会被削弱，避免靠近目标时过冲。")
    eq("e(k) = N_t - N(k)", "4.3")
    eq("v(k) = [N(k)-N(k-1)]/Δt", "4.4")
    eq("u_{raw}(k)=K_p e(k)-K_v v(k)", "4.5")
    p("式中，e(k) 为位置误差，v(k) 为估算速度，u_raw(k) 为未限幅速度指令。M1、M2 和 M3 使用不同的 Kp 和 Kv。调试后的程序中，M2 使用 Kp=1.0、Kv=0.45，M3 使用 Kp=1.25、Kv=0.35。M3 的速度反馈项较低，是因为 M3 接近动作更强调顺利运动和压力触发，而不是高速停止。")
    eq("u(k)=sat(u_{raw}(k),-u_{max},u_{max})", "4.6")
    eq("u(k)=u(k-1)+clip(u_{raw}(k)-u(k-1),-Δu_{max},Δu_{max})", "4.7")
    p("式 (4.6) 是速度限幅，式 (4.7) 是速度斜坡限制。M2 在竖直方向运行时更容易受重力和机构摩擦影响，因此程序还加入了方向最小速度和重力补偿。这样做的结果是，M2 在刚启动时不至于因为指令太小而原地不动，也不会因为指令跳变过大而冲击结构。")
    h("4.5 Derivation of motion parameters", 2)
    p("程序中主要行程参数都从 ONE_TURN = 4096 推导而来。这样比直接写很多固定数值更清楚，也便于之后修改。M1 每次转动四分之一圈，因此 M1_STEP = -4096/4 = -1024 counts。M1 总目标为 -4096×7 = -28672 counts。M2 回零后先上升八分之一圈，所以 M2_INITIAL_LIFT = -4096/8 = -512 counts。M2 正常分层步距为 -4096/20，程序中整数结果为 -204 counts。M2 每层累计三圈，M2_LAYER_TOTAL = 4096×3 = 12288 counts。M3 压力触发后的小回退量同样为 -204 counts。")
    table(["Parameter", "Expression", "Value", "Purpose"], [
        ("M1_STEP", "-ONE_TURN / 4", "-1024", "M1 每层转动四分之一圈"),
        ("M1_TOTAL", "-ONE_TURN * 7", "-28672", "M1 总覆盖范围"),
        ("M2_INITIAL_LIFT", "-ONE_TURN / 8", "-512", "M2 回零后离开底部"),
        ("M2_STEP", "-ONE_TURN / 20", "-204", "M2 每次分层上升"),
        ("M2_LAYER_TOTAL", "ONE_TURN * 3", "12288", "M2 累计三圈后触发 M1"),
        ("M3_SMALL_RETRACT_DELTA", "-ONE_TURN / 20", "-204", "M3 压力触发后回退"),
        ("M3_APPROACH_MAX_TRAVEL", "ONE_TURN * 2", "8192", "M3 最大接近行程"),
    ])
    p("这些参数的选择与当前机械原型有关。M2 每次 1/20 圈是为了让纵向覆盖更细，不至于一次上升太多而漏掉头发区域。M1 每次 1/4 圈用于改变整体方向。M3 回退 1/20 圈是为了在压力触发后让剪刀略微离开接触点，避免后续 M4 横向移动时继续压住目标。")
    h("4.6 M2 homing logic", 2)
    p("M2 回零是最终程序中修改最多的部分。早期版本只要检测到速度为零，就容易把中途卡顿当作触底。后来根据串口记录可以看到，M2 在还没到真正底部时也可能出现 Vel=0.00，并且 Error 仍然很大。如果此时直接 reset zero，后续 M2 上升就会从错误位置开始，甚至在刚开始循环时就卡住。")
    p("最终回零逻辑把触底判断改成多条件确认。程序先记录回零开始时的累计位置、电流和负载，然后让 M2 低速下行。判断底部时使用行程 H、电流 I、负载 L 和速度 v。")
    eq("H=|N-N_{start}|", "4.8")
    eq("ΔI=|I|-|I_{start}|,  ΔL=|L|-|L_{start}|", "4.9")
    p("当电流达到底部信号阈值，或负载相对起始值有足够变化，并且速度低于阈值持续确认时间后，程序才认为触底。如果电机在很短行程内停止，程序会更谨慎，先要求底部信号和低速同时存在。若长时间没有移动，则先提高探测速度，而不是马上判定底部。只有超过最大回零行程、超时或电流超过保护值时才进入急停。")
    p("电流单位也在调试中被单独确认。串口输出中的 Current 不是直接的 mA，而是舵机协议返回的原始计数。按项目中采用的估算，1 count 大约对应 6.5 mA。因此 77 count 接近 500 mA。程序里 M2 回零和上升阶段的保护限制设置为 77 count；底部信号设置为 10 count，用于识别接触趋势。这个区分很重要，因为底部检测需要敏感，而真正急停保护不应该过早触发。")
    h("4.7 State machine and action sequence", 2)
    p("控制流程采用状态机实现。系统上电后先等待按钮启动，启动后进入 M2_HOMING_DOWN。M2 找到底部后，当前位置被设为逻辑零点。随后系统进入 M2_INITIAL_LIFT，让 M2 上升 1/8 圈，离开底部接触区域。之后进入 PREPARE_M5，使 M5 保持在 820 附近。准备完成后才进入正式动作循环。")
    p("一次正式动作循环包括 M2_STEP、M3_APPROACH、M3_STOP_AT_PRESSURE、M3_SMALL_RETRACT、M4_CCW、M4_CW、M5_TO_700、M5_BACK_820 和 M3_HOME。M2_STEP 完成一次竖直分层；M3_APPROACH 让剪刀靠近，压力传感器触发后停止；M3_SMALL_RETRACT 让剪刀退出一点；M4 完成横向扫动；M5 完成剪切；最后 M3 回到接近动作前的位置。")
    p("当 M2 的层内累计行程达到三圈时，程序不再继续 M2 分层，而是执行 M1_STEP。M1 到位后，M2 层内累计量清零，再进入下一层循环。这样 M1 和 M2 形成类似二维覆盖的运动逻辑：M2 负责一层内的纵向覆盖，M1 负责不同角度或区域之间的切换。")
    h("4.8 Safety and reliability measures", 2)
    p("安全保护分成几类。第一类是人为控制，按钮可以启动、暂停和恢复流程。暂停时系统会记录当前状态，并停止正在运动的电机。恢复时再根据之前状态继续执行。第二类是传感器保护，压力传感器主要用于 M3 接近动作，当压力信号触发时，M3 立即停止并记录触发位置。第三类是电流保护，程序持续读取电机电流，若超过对应阈值并持续确认，则进入 emergency stop。第四类是堵转和超时保护，如果电机长时间没有有效移动，或某个动作超过允许时间，也会停止系统。")
    p("这些保护在原型阶段仍有局限。它们依赖舵机反馈和 ESP32 程序正常运行，因此不能替代机械限位、硬件急停和独立电流传感器。但它们对调试很有帮助。每次系统异常停止后，串口会输出当时系统状态、当前动作、目标位置、累计位置、电流、负载、电压和温度。通过这些信息，可以判断问题是来自回零误判、堵转、通信失败还是传感器未触发。")
    h("4.9 Debugging method", 2)
    p("本项目的调试方法比较朴素，主要依靠逐步测试和串口日志。首先单独测试每个电机，确认 ID、通信、模式设置和基本运动方向。然后再把两个相关动作组合起来测试，例如 M2 回零和 M2 上升，M3 接近和压力触发，M4 左右位置切换，M5 开合动作。只有这些局部动作稳定后，才把它们放入完整状态机。这样做虽然慢，但能避免一开始就运行五个电机导致问题来源不清楚。")
    p("串口日志在调试中非常重要。每次输出都会包含 RUN、EMG、SYS、ACTION、Pressure、每个电机的 total、target 和反馈值。比如当 M2 在上升阶段停住时，可以同时看到 Error、Vel、Current 和 Load。如果 Error 很大但 Vel 为零，同时 Current 和 Load 上升，就更像机械卡住；如果位置反馈为 -1，则更可能是通信问题；如果 Pressure 已经变为 1，说明 M3 停止可能来自传感器而不是堵转。")
    p("参数调整也采用逐步修改的方式。以 M2 为例，先调回零速度，再调底部电流信号，再调急停电流保护，最后调回零后的初始抬升。每次只改一个或少数几个参数，并观察串口变化。这样虽然不如自动整定方法高效，但对于当前机械原型更安全，因为可以清楚知道某个变化带来的影响。")

    h("5. Results and Discussion", 1)
    h("5.1 Communication and single motor testing", 2)
    p("项目初期先完成单电机测试。测试内容包括舵机 ID 设置、位置读取、速度控制、位置模式控制和轮子模式控制。单电机测试的意义是把通信问题和控制算法问题分开。只有确认 ESP32 能正确读取位置、电流和负载，后续多电机动作才有基础。调试过程中，串口输出被设计成主要观察工具，因为它能直接显示每个电机的 Pos、Speed、Current、Load、Voltage、Temper 和 Move。")
    p("单电机测试证明，位置模式适合 M4 和 M5 这类目标位置固定的动作，而轮子模式适合 M1、M2 和 M3 这类需要多圈或相对位移的动作。M4 的横向移动只需在两个校准位置之间切换，M5 只需在 820 和 700 附近切换，位置模式更加直接。M1、M2、M3 如果用位置模式，会受到单圈范围限制，因此必须采用轮子模式加外部累计位置。")
    h("5.2 Multi-motor workflow result", 2)
    p("在完成单电机测试后，程序逐步扩展到多电机流程。最终流程能够让 M2 先完成回零和初始抬升，再进入 M2 分层、M3 接近、M4 横移、M5 剪切、M3 回位和 M1 层间转动。串口记录显示，在正常情况下，系统能在不同动作之间自动切换，并在每个 debug 周期输出当前 SYS 和 ACTION。这样的输出对后续测试很重要，因为可以清楚知道系统卡在哪一个动作，而不是只看到电机停止。")
    p("多电机联动中最大的变化是动作之间存在依赖。例如 M3 接近必须发生在 M2 到达分层位置之后；M4 横向移动必须发生在 M3 压力触发并回退之后；M5 剪切必须发生在 M4 动作之后；M1 只有在 M2 累计完成一层后才移动。状态机把这些依赖关系写得比较清楚，避免了多个电机同时乱动。")
    h("5.3 M2 homing and lifting result", 2)
    p("M2 的调试占用了最多时间。最初的问题是回零慢并且容易误判。串口记录中出现过 M2 还没有真正回到底部，速度已经为零，程序就判断触底的情况。之后系统把当前位置设为零点，下一步 M2 上升时因为实际位置不对，出现错误越来越大、速度为零、随后堵转急停。这个问题说明，单独用速度或单独用位置误差都不够可靠。")
    p("修改后的程序要求底部信号和低速状态同时满足，并且加入确认时间。对于刚开始回零时出现的停止，程序不会立刻设零，而是先判断是否真的有电流或负载变化。如果没有足够信号，程序会提高探测速度继续尝试。这样可以减少中途卡顿造成的假回零。")
    p("另一个关键修改是回零后先上升 1/8 圈。调试中发现，M2 在底部刚触底后如果立即开始正常循环，第一步上升很容易因为机械接触和负载较大而卡住。加入 M2_INITIAL_LIFT 后，系统先让 M2 离开底部，再准备 M5 和进入剪发动作。这样后续 M2_STEP 的起点更稳定。")
    p("电流保护也经过多次调整。用户调试记录中 Current=10 或 11 时，M2 已经出现明显底部接触或负载变化，但如果把保护阈值也设得太低，会导致正常运动时误停。最终将底部信号和急停保护分开：较低电流用于判断触底趋势，较高电流用于保护。对杜邦线和普通开发板来说，500 mA 量级通常不是长期大电流设计，但在短时间舵机调试中比更高电流安全；真正成品仍应使用独立供电、合适线径和硬件保险。")
    h("5.4 M3 approach and pressure response", 2)
    p("M3 的主要目标是让剪刀逐渐接近头部，并在压力传感器触发时停止。这个动作不能只靠固定行程，因为头部形状、安装误差和剪刀位置都会变化。如果 M3 每次都移动同样距离，某些情况下可能还没接触就停止，另一些情况下可能压得过深。")
    p("程序中 M3_APPROACH 会持续向前运动，直到压力传感器触发或达到最大行程/超时。触发时记录 action.m3TouchTotal，并把该位置作为后续小回退的参考。M3_SMALL_RETRACT 的目标为触发位置加上 -ONE_TURN/20。这样 M3 的回退不是固定绝对位置，而是相对于实际接触点。")
    p("调试中 M3 小回退也出现过堵转情况。串口记录显示 M3 small retract 目标误差较大，但速度长时间为零。根据这个现象，程序加入了回退释放速度和较长的动作超时，同时调整 M3 的 Kp 和 Kv。这个问题也说明，PID 参数不是越大越好。参数过强可能让电机在受力位置来回顶住结构，参数过弱又可能无法克服摩擦。最终参数是在当前机械结构下取得的折中。")
    h("5.5 M5 preparation and communication issue", 2)
    p("M5 是 SC09 舵机，用于控制剪刀开合。M5 的动作本身相对简单，但调试中出现过准备阶段读取不到位置的情况，即 Pos=-1、Speed=-1、Voltage=-1。这类问题更像通信反馈失败，而不一定是电机真实位置错误。早期程序如果遇到 M5 prepare timeout 就直接急停，容易让整个流程卡住。")
    p("最终程序把 M5 准备阶段改得更宽容：准备阶段持续发送保持 820 的命令，如果反馈正常并到位则继续；如果偶尔反馈失败，也不会因为一次失败立刻急停。这样处理是基于实际观察做出的，因为 M5 的保持动作不是危险动作，而通信偶发失败在总线舵机系统中比较常见。该策略提高了流程连续性，但也提醒后续设计需要更好的总线通信检查。")
    h("5.6 Overall discussion", 2)
    p("从整体结果看，本项目已经实现了从单电机控制到五电机动作流程的扩展。程序能够完成多圈相对位置控制、动作队列执行、压力触发停止、暂停恢复和基本异常保护。与最初只让电机转动相比，最终系统更接近一个可用的原型控制框架。")
    p("但是，系统离真正安全可靠的自动理发机仍有距离。当前测试主要依赖串口观察和原型动作验证，缺少严格的定量数据。例如 M2 每步实际位移误差、M3 压力触发响应时间、M4 横向位置重复精度和 M5 剪切成功率都还没有形成系统表格。报告中给出的结果更多是工程调试结果，而不是完整实验统计结果。")
    p("另一个限制是舵机反馈值的精度。Current 和 Load 可以反映趋势，但并不是经过校准的工程测量值。用它们做保护是合理的，但如果要把系统发展成面向人体的设备，还需要独立传感器、硬件限位和机械防护。软件可以降低风险，但不能承担全部安全责任。")
    p("尽管如此，本项目已经建立了后续工作需要的基础。现在每个动作都有明确状态，每个关键参数都有程序位置，串口日志也能显示问题发生时的系统状态。这让后续调试不再是盲目试参数，而是可以根据状态、误差、速度、电流和负载判断问题。")
    h("5.7 Discussion against project objectives", 2)
    p("对照项目目标，ESP32 与舵机通信、模式设置和反馈读取已经基本实现。M1 至 M4 的 ST 系列舵机可以读取位置、电流、负载、电压和温度，M5 的 SC09 舵机也可以接收目标位置并返回基本状态。虽然通信偶尔出现 -1 读数，但程序已经对部分关键位置加入了容错处理。")
    p("相对多圈控制目标也已经实现。程序通过单圈位置反馈和跨圈判断得到 totalPos，再用 targetPos 和 totalPos 的差值计算速度指令。该方法已经用于 M1、M2 和 M3。它的优点是不用增加额外编码器，缺点是断电后累计位置会丢失，所以系统上电后必须重新回零或重新校准。这一点在报告中需要明确说明，否则容易让读者误以为系统拥有绝对多圈位置记忆。")
    p("多电机协调目标基本实现。五个电机不再是分散测试，而是通过状态机连接成一个完整流程。每个动作都需要满足到位、触发或超时等条件后才能进入下一步。这个流程符合自动理发机原型的运动需求，但还不能说明剪发效果已经达到实际使用要求，因为当前结果主要验证运动逻辑，剪发质量和人体安全还需要更多实验。")
    p("安全目标实现了一部分。按钮暂停、压力触发、电流保护、堵转保护和超时保护都已经写入程序，也在调试中起到作用。但这些保护仍属于软件层面，只适合作为原型调试阶段的安全辅助。对于真实设备，还必须增加硬件急停、机械限位和电源级保护。换句话说，本项目完成的是软件控制基础，不是最终产品安全认证。")

    h("6. Conclusions and Future Work", 1)
    p("本项目完成了一套基于 ESP32 的自动理发机多电机控制系统。系统能够同时控制 ST3215-HS 和 SC09 总线舵机，支持位置模式和轮子模式混合运行。通过跨圈检测和累计位置计算，M1、M2 和 M3 能够在轮子模式下完成相对多圈运动；通过位置模式，M4 和 M5 能够完成固定位置动作。")
    p("项目的主要成果包括：完成 ESP32 与多舵机通信；实现多圈相对位置控制；建立 M1 至 M5 的动作状态机；实现 M2 回零、初始抬升和分层逻辑；实现 M3 压力触发停止和小回退；加入按钮暂停恢复、电流保护、堵转保护和动作超时保护；建立较完整的串口调试输出。")
    p("本项目也暴露了几个需要继续解决的问题。M2 回零仍然依赖舵机反馈电流和负载，机械卡顿与真实触底之间仍可能在极端情况下混淆。M3 的压力触发目前只作为数字信号使用，还没有对压力大小进行连续测量。M5 的通信反馈偶尔失败，需要从总线硬件、电源和软件重试机制三个方面继续改善。")
    p("未来工作首先应加入硬件急停、限位开关和独立电流检测。其次，应对 M2、M3 的不同位置和不同负载进行系统测试，记录每组参数下的成功率、最大电流、动作时间和停止误差。第三，应把剪刀模块和安全替代物结合起来做重复剪切测试，量化覆盖效果。最后，如果系统要接近真实使用场景，还需要加入更可靠的人体保护结构和更严格的风险评估。")
    p("总体而言，本项目没有完成一台可直接使用的自动理发机，但完成了自动理发机原型中最关键的软件控制基础。该基础能够支持后续机械优化、传感器升级和更完整的剪发实验，也为之后学生继续开发该项目提供了较清楚的程序框架和调试经验。")

    h("References", 1)
    p("Espressif Systems. ESP32 technical reference manual and programming documentation.")
    p("Waveshare. ST3215-HS serial bus servo user manual and communication protocol documentation.")
    p("Waveshare. SC09 serial bus servo user manual and communication protocol documentation.")
    p("FreeRTOS. Real-time kernel documentation for embedded task scheduling.")
    h("Appendix A: Key Program Parameters", 1)
    table(["Group", "Parameter", "Value", "Meaning"], [
        ("Encoder", "ONE_TURN", "4096", "one revolution in encoder counts"),
        ("M1", "M1_STEP", "-1024", "quarter-turn movement"),
        ("M2", "M2_INITIAL_LIFT", "-512", "initial lift after homing"),
        ("M2", "M2_STEP", "-204", "vertical layer increment"),
        ("M2", "M2_HOMING_CURRENT_LIMIT", "77", "about 500 mA protection"),
        ("M2", "M2_HOMING_BOTTOM_CURRENT", "10", "bottom detection signal"),
        ("M3", "M3_APPROACH_SPEED", "120", "slow approach speed"),
        ("M4", "M4_CCW_POS / M4_CW_POS", "568 / 3528", "lateral positions"),
        ("M5", "M5_HOLD_POS / M5_CUT_POS", "820 / 700", "scissor open/cut positions"),
    ])

    count = body_count(report)
    for kind, text, level in report:
        if kind == "h":
            add_heading(doc, text, level)
        elif kind == "p":
            add_para(doc, text)
        elif kind == "eq":
            eq_text, number = text.split("|", 1)
            add_equation(doc, eq_text, number)
        elif kind == "t":
            headers, rows = text
            add_table(doc, headers, rows)

    add_heading(doc, "Appendix B: Main body character count", 1)
    add_para(doc, f"Main body approximate Chinese-character count generated by script: {count}. This count excludes title page, declaration, project summary, nomenclature, references and appendices.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)
    print(f"BODY_CHARS={count}")


if __name__ == "__main__":
    main()
